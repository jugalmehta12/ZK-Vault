from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def _set_run_font(run, size=12, bold=False):
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)


def create_report():
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)

    # Header (top left)
    h = section.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h_run = h.add_run('DEPSTAR/CSE/2026/Roll No.')
    _set_run_font(h_run, size=10, bold=False)

    # Footer (left-center-right)
    footer_table = section.footer.add_table(rows=1, cols=3, width=Inches(6.5))
    footer_table.autofit = True

    p_left = footer_table.rows[0].cells[0].paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_left = p_left.add_run('DEPSTAR')
    _set_run_font(r_left, size=10, bold=False)

    p_mid = footer_table.rows[0].cells[1].paragraphs[0]
    p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(p_mid)

    p_right = footer_table.rows[0].cells[2].paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_right = p_right.add_run('Department of Computer Science and Engineering')
    _set_run_font(r_right, size=10, bold=False)

    def chapter_heading(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text.upper())
        _set_run_font(r, size=16, bold=True)
        p.paragraph_format.line_spacing = 1.5

    def section_heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text.upper())
        _set_run_font(r, size=14, bold=True)
        p.paragraph_format.line_spacing = 1.5

    def subsection_heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        _set_run_font(r, size=12, bold=True)
        p.paragraph_format.line_spacing = 1.5

    def body(text, center=False):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5

    def bullet(text):
        p = doc.add_paragraph(text, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5

    def mermaid_figure(caption, code):
        block = doc.add_paragraph()
        block.alignment = WD_ALIGN_PARAGRAPH.LEFT
        block.paragraph_format.line_spacing = 1.0
        run = block.add_run('```mermaid\n' + code.strip() + '\n```')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing = 1.5
        cap_run = cap.runs[0]
        cap_run.font.name = 'Times New Roman'
        cap_run.font.size = Pt(12)
        cap_run.bold = True

    def ascii_figure(caption, art):
        block = doc.add_paragraph()
        block.alignment = WD_ALIGN_PARAGRAPH.LEFT
        block.paragraph_format.line_spacing = 1.0
        run = block.add_run('```text\n' + art.strip() + '\n```')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing = 1.5
        cap_run = cap.runs[0]
        cap_run.font.name = 'Times New Roman'
        cap_run.font.size = Pt(12)
        cap_run.bold = True

    # COVER PAGE
    chapter_heading('PROJECT REPORT')
    body('', center=True)
    body('', center=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ZERO-KNOWLEDGE PASSWORD VAULT WITH CLIENT-SIDE RISK INTELLIGENCE')
    _set_run_font(r, size=16, bold=True)
    p.paragraph_format.line_spacing = 1.5

    body('', center=True)
    body('Submitted in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering', center=True)
    body('', center=True)
    body('By', center=True)
    body('Jugal Mehta (24DCS055)', center=True)
    body('Adarsh Nanera (24DCS058)', center=True)
    body('Yug Patel (24DCS097)', center=True)
    body('', center=True)
    body('Under the Guidance of', center=True)
    body('Asst. Prof. Naina Parmar', center=True)
    body('', center=True)
    body('DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING', center=True)
    body('DEPSTAR, FACULTY OF TECHNOLOGY AND ENGINEERING', center=True)
    body('CHARUSAT, CHANGA', center=True)
    body('APRIL 2026', center=True)
    doc.add_page_break()

    # CANDIDATE CERTIFICATE
    chapter_heading('CANDIDATE CERTIFICATE')
    body('This is to certify that the project report entitled "Zero-Knowledge Password Vault with Client-Side Risk Intelligence" submitted by Jugal Mehta (24DCS055), Adarsh Nanera (24DCS058), and Yug Patel (24DCS097) in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering at DEPSTAR, Faculty of Technology and Engineering, CHARUSAT is a bona fide work carried out by them under my guidance and supervision during the academic year 2025-26.')
    body('To the best of my knowledge, this work has not been submitted in full or in part to this or any other University/Institute for the award of any degree or diploma.')
    body('')
    body('Guide Signature: ____________________')
    body('Name: Asst. Prof. Naina Parmar')
    body('Department of Computer Science and Engineering, DEPSTAR')
    body('Date: ____________________')
    doc.add_page_break()

    # CANDIDATE DECLARATION
    chapter_heading('CANDIDATE DECLARATION')
    body('We hereby declare that the project report titled "Zero-Knowledge Password Vault with Client-Side Risk Intelligence" is our original work carried out at DEPSTAR, Faculty of Technology and Engineering, CHARUSAT under the guidance of Asst. Prof. Naina Parmar. The report is being submitted in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering.')
    body('We further declare that this work has not been submitted, either in part or full, to any other institution or university for the award of any degree or diploma.')
    body('')
    body('Jugal Mehta (24DCS055)   Signature: ____________________')
    body('Adarsh Nanera (24DCS058) Signature: ____________________')
    body('Yug Patel (24DCS097)     Signature: ____________________')
    body('Date: ____________________')
    doc.add_page_break()

    # TABLE OF CONTENTS
    chapter_heading('TABLE OF CONTENTS')
    body('Cover Page.................................................................................i', center=False)
    body('Candidate Certificate....................................................................ii', center=False)
    body('Candidate Declaration...................................................................iii', center=False)
    body('Abstract........................................................................................iv', center=False)
    body('Chapter 1: Introduction....................................................................1', center=False)
    body('Chapter 2: Literature Review............................................................6', center=False)
    body('Chapter 3: System Analysis.................................................................11', center=False)
    body('Chapter 4: Technology Stack..............................................................17', center=False)
    body('Chapter 5: System Design..................................................................23', center=False)
    body('Chapter 6: Testing.............................................................................33', center=False)
    body('Chapter 7: Results..............................................................................39', center=False)
    body('Chapter 8: Challenges Faced...............................................................43', center=False)
    body('Chapter 9: Conclusion and Future Scope.............................................47', center=False)
    body('References....................................................................................51', center=False)
    body('Appendices....................................................................................53', center=False)
    doc.add_page_break()

    # ABSTRACT
    chapter_heading('ABSTRACT')
    body('Password reuse, weak password selection, and insecure storage practices remain key causes of account compromise in modern web systems. This project presents a Zero-Knowledge Password Vault with Client-Side Risk Intelligence that addresses these issues through privacy-preserving cryptography and actionable security analytics. The implemented system follows a strict separation of responsibilities: the browser performs all sensitive cryptographic operations, while the server acts only as an authenticated ciphertext storage and session management service. A master password is never transmitted to the server. Instead, the client derives a non-extractable AES-GCM key using PBKDF2 with SHA-256 and 310,000 iterations, encrypts the vault locally, and sends only encrypted payloads to the backend. The backend stores user account metadata, bcrypt-hashed login password, and encrypted vault blobs in MongoDB, with additional controls such as rate limiting, CSRF protection, input validation, secure cookies, and audit logging. The system further integrates client-side risk intelligence including entropy-based strength scoring, duplicate password detection, stale password identification, and breach exposure checks using the Have I Been Pwned k-anonymity protocol. The developed prototype demonstrates that strong usability and strong privacy are not mutually exclusive. Experimental usage confirms practical response times, successful end-to-end encrypted vault synchronization, and meaningful security insights without disclosure of plaintext credentials to the server.')
    doc.add_page_break()

    # CHAPTER 1
    chapter_heading('CHAPTER 1: INTRODUCTION')
    section_heading('1.1 Background')
    body('The rapid growth of digital services has significantly increased the number of credentials that a typical user must manage. In practice, many users adopt insecure habits such as storing passwords in plain notes, using predictable patterns, or reusing the same password across multiple platforms. These habits create a high-impact risk because a single breach can cascade into compromise of multiple accounts. Password managers were introduced to solve this challenge; however, trust in the service provider remains a concern when server-side infrastructure has access to decryptable data.')
    section_heading('1.2 Problem Definition')
    body('The core problem addressed in this project is the lack of trustworthy, privacy-preserving password storage with proactive risk awareness. Traditional architectures may encrypt data at rest but still allow privileged server components to access plaintext under certain conditions. This creates a trust bottleneck and a single point of failure. The project therefore targets a model where the server can authenticate users and store data but can never decrypt user vault content.')
    section_heading('1.3 Objectives')
    bullet('To implement a web-based password vault that follows strict zero-knowledge principles.')
    bullet('To perform encryption and decryption exclusively in the browser using Web Crypto API.')
    bullet('To use robust cryptographic primitives: PBKDF2-SHA256 for key derivation and AES-GCM-256 for authenticated encryption.')
    bullet('To provide secure account authentication with optional TOTP-based multi-factor authentication.')
    bullet('To implement client-side password risk intelligence including strength, breach, duplicate, and age analysis.')
    bullet('To maintain secure operational controls including CSRF defense, rate limiting, validation, and audit logging.')
    section_heading('1.4 Scope')
    body('The current scope includes account registration and login, encrypted vault CRUD operations, risk scoring dashboard, session/device visibility, and vault version rollback. The system is designed as a production-oriented academic prototype and is suitable for extension into enterprise or mobile use cases. Browser extension autofill, hardware key integration, and distributed encrypted storage are intentionally left for future work.')
    section_heading('1.5 Motivation')
    body('The project was motivated by two practical observations: first, users require simple workflows for secure password management; second, a privacy-first architecture should not depend on blind trust in backend operators. By combining zero-knowledge storage with actionable risk intelligence, the system not only protects secret data but also improves day-to-day password hygiene.')
    doc.add_page_break()

    # CHAPTER 2
    chapter_heading('CHAPTER 2: LITERATURE REVIEW')
    section_heading('2.1 Overview of Existing Password Managers')
    body('Modern password managers generally provide secure storage, generator utilities, and cross-device synchronization. Open-source solutions such as Bitwarden emphasize encrypted vault storage and transparency, while proprietary cloud-based solutions focus on convenience and ecosystem integration. Browser-native managers offer ease of use but may not provide comprehensive risk intelligence or architecture-level guarantees on key handling.')
    section_heading('2.2 Research on Zero-Knowledge Security')
    body('Zero-knowledge password management designs separate user authentication from vault decryption capability. Prior architectural references highlight that secrecy depends not only on cipher strength but also on key custody. If key material can be reconstructed server-side, confidentiality assumptions weaken. Accordingly, client-only key derivation and non-extractable key usage represent a stronger practical model for end-user privacy.')
    section_heading('2.3 Password Security and Breach Intelligence')
    body('NIST SP 800-63B and OWASP recommendations emphasize strong password policies, avoidance of predictable patterns, and screening against breached credential datasets. The Have I Been Pwned k-anonymity method demonstrates how breach checking can be performed with reduced disclosure by querying only hash prefixes. This approach aligns with privacy-by-design principles and is appropriate for client-side validation flows.')
    section_heading('2.4 Comparative Analysis')
    body('Compared with many baseline managers, the proposed system contributes a project-specific combination: strict client-side cryptography, privacy-preserving breach checks, vault-level intelligence scoring, and backend hardening controls mapped to common web attack vectors. This integrated design supports both confidentiality and practical maintainability in an educational full-stack environment.')
    doc.add_page_break()

    # CHAPTER 3
    chapter_heading('CHAPTER 3: SYSTEM ANALYSIS')
    section_heading('3.1 Functional Requirements')
    bullet('User registration, login, logout, and profile session continuity.')
    bullet('Encrypted vault retrieval and update through authenticated API routes.')
    bullet('Client-side decryption after successful login and master password entry.')
    bullet('CRUD operations for vault entries with search and category filtering.')
    bullet('MFA setup, verification, and disable flows using TOTP.')
    bullet('Vault version history and rollback support.')
    bullet('Password generation and real-time risk analysis dashboard.')
    bullet('Session listing and selective session revocation.')
    section_heading('3.2 Non-Functional Requirements')
    bullet('Security: resistance to brute force, CSRF, injection, and session hijacking.')
    bullet('Confidentiality: server must not access plaintext vault or master password.')
    bullet('Availability: responsive APIs and graceful behavior under failed external checks.')
    bullet('Performance: acceptable key-derivation delay with strong security parameters.')
    bullet('Maintainability: modular frontend JS and route-wise backend separation.')
    bullet('Scalability: session and audit model extensible for distributed deployment.')
    section_heading('3.3 Existing vs Proposed System')
    body('In conventional systems, vault encryption may occur server-side or through models where backend services can access decryptable states. In the proposed architecture, encryption takes place before transmission and the backend stores only opaque ciphertext with metadata (salt and IV). This materially reduces breach impact because exposure of database records does not provide plaintext credential access.')
    section_heading('3.4 Feasibility')
    subsection_heading('3.4.1 Technical Feasibility')
    body('The stack relies on mature, well-documented technologies: Node.js, Express, MongoDB, and Web Crypto API. The selected primitives are supported in modern browsers and integrate directly into a modular JavaScript frontend.')
    subsection_heading('3.4.2 Economic Feasibility')
    body('The system uses open-source tools and local development infrastructure, making implementation economically feasible in an academic setting.')
    subsection_heading('3.4.3 Operational Feasibility')
    body('The user interface supports straightforward workflows for account operations and vault management. Security checks are embedded into normal user interactions rather than requiring specialist knowledge.')
    doc.add_page_break()

    # CHAPTER 4
    chapter_heading('CHAPTER 4: TECHNOLOGY STACK')
    section_heading('4.1 Frontend Technologies')
    body('The frontend is implemented in HTML, CSS, and modular Vanilla JavaScript to minimize dependency risk and preserve full control over sensitive memory handling. Cryptographic operations are performed through Web Crypto API primitives: PBKDF2 key derivation with SHA-256, AES-GCM encryption/decryption, and SHA-1 digesting for breach lookup prefixing. The UI includes dedicated modules for authentication, vault orchestration, cryptography, risk intelligence, and password generation.')
    section_heading('4.2 Backend Technologies')
    body('The backend uses Node.js with Express for routing and middleware composition. MongoDB with Mongoose handles persistence for user accounts, sessions, audit logs, and vault versions. Authentication uses JWT in secure cookie transport. Auxiliary security libraries include Helmet, express-rate-limit, express-validator, express-mongo-sanitize, cookie-parser, and CORS controls. TOTP MFA support is provided by speakeasy and QR code generation by qrcode.')
    section_heading('4.3 Database Layer')
    body('MongoDB collections are structured for operational clarity: User stores account identity and encrypted vault blob; Session stores per-device session state; AuditLog stores security-relevant events with metadata; VaultVersion stores encrypted snapshots for rollback. This design supports traceability and controlled recovery without violating zero-knowledge constraints.')
    section_heading('4.4 Why These Technologies Were Chosen')
    bullet('Web Crypto API provides standards-based browser-native cryptography and non-extractable key options.')
    bullet('Express enables modular middleware security composition and maintainable API routing.')
    bullet('MongoDB supports flexible document schemas and rapid iteration for project-scale prototypes.')
    bullet('JWT with secure cookies supports stateless verification with controlled session tracking in database.')
    bullet('Jest and Supertest provide effective backend test automation for route behavior and security checks.')
    doc.add_page_break()

    # CHAPTER 5
    chapter_heading('CHAPTER 5: SYSTEM DESIGN')
    section_heading('5.1 Use Case Description')
    bullet('Actor: End User. Use cases: register account, login, unlock vault, manage passwords, analyze risk, configure MFA, revoke sessions, rollback vault.')
    bullet('Actor: System. Use cases: validate requests, issue and verify session tokens, persist encrypted vault payloads, generate audit events.')
    section_heading('5.2 High-Level Architecture')
    body('The design follows a browser-centric trust boundary. The frontend performs all sensitive cryptographic transformations and sends only ciphertext to backend endpoints. The backend verifies identity and authorizes requests but never receives the master secret. Communication is via HTTPS with cookie-based session flow and CSRF token checks.')

    architecture_diagram = """
flowchart TB
    U[User]

    subgraph CLIENT[Client Tier]
        FE[React Frontend UI]
        CE[Client-Side Cryptography\\nWeb Crypto API\\nPBKDF2 + AES-GCM]
    end

    subgraph SERVER[Application Tier]
        API[Express and Node.js REST API]
    end

    subgraph DATA[Data Tier]
        DB[(MongoDB Encrypted Vault Store)]
    end

    U -->|User actions| FE
    FE -->|Master password input retained locally| CE
    CE -->|Encrypt vault then send ciphertext, IV, salt| API
    API -->|Persist encrypted blob| DB

    DB -->|Fetch encrypted vault blob| API
    API -->|Encrypted API response| FE
    FE -->|Decrypt using derived key| CE
    CE -->|Decrypted vault in browser memory| FE

    API -.->|Zero-knowledge boundary server never receives plaintext| CE
"""
    mermaid_figure('Fig 5.1 System Architecture Diagram', architecture_diagram)

    use_case_diagram = """
flowchart LR
    A[Actor: User]

    subgraph SYS[System Boundary: Zero-Knowledge Password Vault]
        UC1((Register Account))
        UC2((Login))
        UC3((Add Credential))
        UC4((View Credential))
        UC5((Edit Credential))
        UC6((Delete Credential))
        UC7((Generate Strong Password))
        UC8((Run Password Risk Analysis))
        UC9((Logout))
    end

    A --- UC1
    A --- UC2
    A --- UC3
    A --- UC4
    A --- UC5
    A --- UC6
    A --- UC7
    A --- UC8
    A --- UC9
"""
    mermaid_figure('Fig 5.2 Use Case Diagram', use_case_diagram)

    er_diagram = """
erDiagram
    USERS {
        string user_id PK
        string email UK
        string master_password_hash
    }

    VAULT_ENTRIES {
        string entry_id PK
        string website
        string username
        string encrypted_password
        string user_id FK
    }

    RISK_ANALYSIS {
        string analysis_id PK
        int strength_score
        boolean reuse_detected
        string user_id FK
    }

    USERS ||--o{ VAULT_ENTRIES : stores
    USERS ||--o{ RISK_ANALYSIS : generates
"""
    mermaid_figure('Fig 5.3 ER Diagram', er_diagram)

    dfd_level0 = """
flowchart LR
    E1[External Entity: User]
    P0((P0: Password Vault System))
    D1[(D1: Encrypted Storage)]

    E1 -->|Authentication data and credential actions| P0
    P0 -->|Encrypted vault records| D1
    D1 -->|Encrypted vault records| P0
    P0 -->|Vault view and security insights| E1
"""
    mermaid_figure('Fig 5.4 Data Flow Diagram Level 0', dfd_level0)

    dfd_level1 = """
flowchart LR
    E1[External Entity: User]
    P1((1.0 Authentication Process))
    P2((2.0 Encryption and Decryption Module))
    P3((3.0 Vault Management Module))
    P4((4.0 Risk Intelligence Engine))
    D1[(D1: User and Session Store)]
    D2[(D2: Encrypted Vault Store)]

    E1 -->|Register and login data| P1
    P1 -->|Authentication status| E1
    P1 <--> |User profile and session data| D1

    E1 -->|Create, read, update, delete requests| P3
    P3 -->|Vault operation results| E1

    P3 -->|Vault payload for encryption| P2
    P2 -->|Ciphertext + IV + Salt| P3
    P3 <--> |Encrypted vault records| D2

    E1 -->|Risk scan request| P4
    P3 -->|In-memory vault entries| P4
    P4 -->|Strength, reuse, breach, age metrics| E1
"""
    mermaid_figure('Fig 5.5 Data Flow Diagram Level 1', dfd_level1)
    section_heading('5.3 Authentication Flow')
    body('During registration, user account password is validated and hashed with bcrypt before storage. During login, credentials are verified and JWT is issued with a generated session identifier. If MFA is enabled, TOTP validation is required before token issuance. A valid authenticated state then permits retrieval of encrypted vault data. Master password handling remains exclusively client-side.')

    login_sequence = """
sequenceDiagram
    autonumber
    actor User
    participant FE as Frontend
    participant CR as Crypto Layer
    participant BE as Backend API
    participant DB as Database

    User->>FE: Enter account email and password
    FE->>BE: POST /api/auth/login
    BE->>DB: Query user by email
    DB-->>BE: User record and bcrypt hash
    BE->>BE: Verify password hash
    alt MFA Enabled
        BE-->>FE: MFA required
        FE->>BE: Submit TOTP token
        BE->>BE: Validate TOTP
    end
    BE-->>FE: Login success and secure session cookie

    User->>FE: Enter master password
    FE->>BE: GET /api/vault
    BE->>DB: Read encryptedVault, vaultIV, vaultSalt
    DB-->>BE: Encrypted vault package
    BE-->>FE: Return encrypted payload
    FE->>CR: Derive key using PBKDF2
    CR-->>FE: Derived AES-GCM key (non-extractable)
    FE->>CR: Decrypt vault package
    CR-->>FE: Plaintext vault entries
    FE-->>User: Render unlocked dashboard
"""
    mermaid_figure('Fig 5.6 Sequence Diagram for User Login', login_sequence)
    section_heading('5.4 Encryption and Decryption Workflow')
    body('When unlocking vault, the browser derives an AES key from master password and stored salt using PBKDF2 (310000 iterations). On save, vault JSON is encrypted with AES-GCM and a fresh IV, then uploaded as base64 ciphertext. On load, ciphertext is downloaded and decrypted locally using derived key and IV. Authentication tags in AES-GCM ensure tamper detection during decryption.')

    add_credential_sequence = """
sequenceDiagram
    autonumber
    actor User
    participant FE as Frontend
    participant CR as Crypto Layer
    participant BE as Backend API
    participant DB as Database

    User->>FE: Create new credential
    FE->>FE: Validate entry fields
    FE->>FE: Update in-memory vault model
    FE->>CR: Encrypt updated vault using AES-GCM
    CR-->>FE: Ciphertext package (ciphertext, IV, salt)
    FE->>BE: PUT /api/vault with encrypted package
    BE->>DB: Store encrypted vault blob
    DB-->>BE: Persist success
    BE-->>FE: Save confirmation
    FE-->>User: Display updated vault list
"""
    mermaid_figure('Fig 5.7 Sequence Diagram for Adding New Credential', add_credential_sequence)

    zk_encryption_workflow = """
flowchart TD
    MP[Master Password]
    KDF[PBKDF2 Key Derivation]
    ENC[AES-GCM Encryption]
    STORE[(Ciphertext Storage in MongoDB)]
    RET[Retrieve Ciphertext]
    DEC[Client-side Decryption]

    MP --> KDF
    KDF --> ENC
    ENC --> STORE
    STORE --> RET
    RET --> DEC
"""
    mermaid_figure('Fig 5.8 Zero-Knowledge Encryption Workflow Diagram', zk_encryption_workflow)
    section_heading('5.5 ER Diagram Explanation')
    body('Core entities are User, Session, AuditLog, and VaultVersion. A User has many Sessions, many AuditLog records, and many VaultVersion records. Vault entry objects are nested within encrypted payload and therefore are not relationally visible server-side. This intentionally preserves confidentiality at storage level.')

    activity_diagram = """
flowchart TD
    A([Start]) --> B[Launch application]
    B --> C[Submit login credentials]
    C --> D{Credentials valid?}
    D -->|No| E[Display authentication error]
    E --> C
    D -->|Yes| F[Enter master password]
    F --> G[Derive key via PBKDF2]
    G --> H[Retrieve encrypted vault]
    H --> I[Decrypt vault locally using AES-GCM]
    I --> J[Display vault dashboard]

    J --> K[Perform vault operations]
    K --> L[Run password risk analysis]
    L --> M[Encrypt updated vault]
    M --> N[Store encrypted payload]
    N --> O([End])
"""
    mermaid_figure('Fig 5.9 Activity Diagram', activity_diagram)
    section_heading('5.6 Module Description')
    bullet('Auth Module: registration, login, MFA setup/verify/disable, logout, session listing and revocation.')
    bullet('Vault Module: encrypted blob retrieval, encrypted blob update, version listing, rollback.')
    bullet('Crypto Module: key derivation, encryption/decryption, binary/base64 conversion utilities.')
    bullet('Risk Intelligence Module: strength score, breach check via k-anonymity, duplicate and stale detection.')
    bullet('Generator Module: cryptographically secure password generation with entropy estimation.')

    module_hierarchy = """
flowchart TB
    ROOT[Password Vault System]
    AUTH[Authentication Module]
    VM[Vault Management Module]
    ENC[Encryption Module]
    RI[Risk Intelligence Module]
    DM[Database Module]

    ROOT --> AUTH
    ROOT --> VM
    ROOT --> ENC
    ROOT --> RI
    ROOT --> DM

    AUTH --> A1[Registration and Login]
    AUTH --> A2[MFA and Session Governance]

    VM --> V1[Credential CRUD]
    VM --> V2[Versioning and Rollback]

    ENC --> E1[PBKDF2 Key Derivation]
    ENC --> E2[AES-GCM Encryption and Decryption]

    RI --> R1[Strength Scoring]
    RI --> R2[Reuse, Breach, and Age Analysis]

    DM --> D1[Users and Sessions]
    DM --> D2[Encrypted Vault and Audit Logs]
"""
    mermaid_figure('Fig 5.10 Module Hierarchy Diagram', module_hierarchy)
    section_heading('5.7 Data Flow Explanation')
    body('Sensitive data enters at the browser UI, is transformed into encrypted payloads, then transmitted to backend APIs and persisted in MongoDB. Reverse flow retrieves encrypted data to the browser where local decryption restores runtime vault state. At no stage is plaintext vault content transported or stored outside the user runtime context.')
    doc.add_page_break()

    # CHAPTER 6
    chapter_heading('CHAPTER 6: TESTING')
    section_heading('6.1 Testing Strategy')
    body('Testing was conducted at unit, integration, and system levels. Backend API behavior was validated through Jest and Supertest. Cryptographic flow checks ensured that only encrypted payloads are persisted. Security controls such as rate limiting and CSRF validation were tested under expected and adverse request patterns.')
    section_heading('6.2 Unit Testing')
    bullet('Auth route behavior: register success/failure, login invalid credentials, MFA-required response.')
    bullet('Vault route behavior: encrypted payload validation, unauthorized access rejection.')
    bullet('Session validation middleware: token/session mismatch and expiry handling.')
    section_heading('6.3 Integration Testing')
    bullet('End-to-end auth to vault flow with cookie session propagation.')
    bullet('Encrypt-save-fetch-decrypt cycle with expected plaintext equivalence.')
    bullet('Tamper handling by modifying ciphertext/IV and verifying decryption failure.')
    section_heading('6.4 System Testing')
    bullet('Browser workflow: register, login, unlock, add/edit/delete entries, logout.')
    bullet('Risk dashboard behavior with weak, reused, breached, and old passwords.')
    bullet('Session controls and audit log visibility after authentication events.')

    testing_flow = """
flowchart LR
    IN[Input Data]
    VAL[Validation Stage]
    DEC{Validation Passed?}
    ENC[Encryption Stage]
    STG[Encrypted Storage Stage]
    RSK[Risk Analysis Stage]
    OUT[Output and Reporting]

    IN --> VAL
    VAL --> DEC
    DEC -->|Yes| ENC
    DEC -->|No| OUT
    ENC --> STG
    STG --> RSK
    RSK --> OUT
"""
    mermaid_figure('Fig 6.1 Testing Flow Diagram', testing_flow)
    section_heading('6.5 Test Case Table')
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'TC ID'
    hdr[1].text = 'Test Scenario'
    hdr[2].text = 'Input/Precondition'
    hdr[3].text = 'Expected Output'
    hdr[4].text = 'Actual Output'
    hdr[5].text = 'Status'
    test_rows = [
        ('TC-01', 'User Registration', 'Valid email and strong password', 'Account created; JWT cookie set', 'As expected', 'Pass'),
        ('TC-02', 'Login Failure', 'Invalid password', '401 Unauthorized', 'As expected', 'Pass'),
        ('TC-03', 'MFA Required', 'MFA-enabled account without token', 'mfaRequired true response', 'As expected', 'Pass'),
        ('TC-04', 'Vault Save', 'Encrypted blob, IV, salt provided', 'Vault stored successfully', 'As expected', 'Pass'),
        ('TC-05', 'Vault Read', 'Authenticated user', 'Encrypted vault payload returned', 'As expected', 'Pass'),
        ('TC-06', 'Rate Limit', 'Repeated login attempts > threshold', '429 Too Many Requests', 'As expected', 'Pass'),
        ('TC-07', 'Breach Check', 'Common leaked password', 'Breached warning shown', 'As expected', 'Pass'),
        ('TC-08', 'Duplicate Detection', 'Two entries same password', 'Duplicate risk flagged', 'As expected', 'Pass'),
    ]
    for row in test_rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    doc.add_page_break()

    # CHAPTER 7
    chapter_heading('CHAPTER 7: RESULTS')
    section_heading('7.1 Functional Outcomes')
    body('The implemented system supports the complete planned workflow: secure registration and login, optional MFA, encrypted vault management, session controls, risk intelligence analytics, and rollback-capable vault versioning. All core endpoints and client modules interacted consistently in integration testing.')
    section_heading('7.2 Security Outcomes')
    body('Observed network payloads during vault synchronization contained only encrypted base64 blobs with IV and salt metadata. MongoDB persisted ciphertext-only vault content. Brute-force mitigation and CSRF controls behaved as expected. Session invalidation and expiry checks enforced re-authentication where appropriate.')
    section_heading('7.3 Performance Observations')
    body('PBKDF2 derivation introduced a deliberate client-side delay that strengthens offline attack resistance. AES-GCM encryption/decryption for normal vault sizes remained practically instantaneous. API response times were suitable for interactive usage in local deployment conditions.')
    section_heading('7.4 Interface Snapshots (ASCII Wireframes)')

    login_wireframe = """
+--------------------------------------------------------------------+
| ZK Vault                                                           |
|--------------------------------------------------------------------|
| Email Address          [ you@example.com                      ]    |
| Account Password       [ ***************                      ]    |
| Authenticator Code     [ 123456 (optional if MFA enabled)     ]    |
|                                                                    |
| [ Sign In ]                                     [ Create Account ] |
|                                                                    |
| Security Note: Master password is never sent to server.            |
+--------------------------------------------------------------------+
"""
    ascii_figure('Fig 7.1 Login Page Wireframe', login_wireframe)

    dashboard_wireframe = """
+--------------------------------------------------------------------------------+
| ZK Vault Dashboard                                      user@example.com      |
|--------------------------------------------------------------------------------|
| Search [ gmail ]   Category [ All ]    [ + Add Credential ] [ Generate ]      |
|--------------------------------------------------------------------------------|
| Title      Username             URL                      Risk      Last Update  |
| Gmail      jugal.mail           mail.google.com          Medium    2026-04-18   |
| GitHub     jugal-dev            github.com               Strong    2026-04-15   |
| Netflix    jugal.stream         netflix.com              Weak      2026-01-04   |
|--------------------------------------------------------------------------------|
| [ View ] [ Edit ] [ Delete ] [ Save Vault ] [ History ] [ Logout ]            |
+--------------------------------------------------------------------------------+
"""
    ascii_figure('Fig 7.2 Vault Dashboard Wireframe', dashboard_wireframe)

    risk_wireframe = """
+--------------------------------------------------------------------------+
| Risk Intelligence Panel                                                  |
|--------------------------------------------------------------------------|
| Overall Security Score: 72 / 100   Status: FAIR                         |
|                                                                          |
| Strength Summary:                                                        |
|  - Very Strong: 04                                                       |
|  - Strong:      06                                                       |
|  - Weak:        02                                                       |
|                                                                          |
| Alerts:                                                                  |
|  - Reused Passwords: 1                                                   |
|  - Breached Passwords (HIBP): 1                                          |
|  - Old Passwords (>90 days): 3                                           |
|                                                                          |
| Recommended Actions: Update weak/reused credentials immediately.         |
+--------------------------------------------------------------------------+
"""
    ascii_figure('Fig 7.3 Risk Analysis Wireframe', risk_wireframe)

    mfa_wireframe = """
+--------------------------------------------------------------------+
| Multi-Factor Authentication Setup                                   |
|--------------------------------------------------------------------|
| Step 1: Scan QR in Authenticator App                               |
| [ QR CODE AREA ]                                                   |
|                                                                    |
| Step 2: Enter 6-digit token                                        |
| Verification Token      [ 654321 ]                                 |
|                                                                    |
| [ Verify & Enable MFA ]   [ Disable MFA ]                          |
+--------------------------------------------------------------------+
"""
    ascii_figure('Fig 7.4 MFA Setup Wireframe', mfa_wireframe)

    session_wireframe = """
+--------------------------------------------------------------------------------+
| Session Management                                                             |
|--------------------------------------------------------------------------------|
| Session ID      Device/Browser        IP Address       Last Seen       Status  |
|--------------------------------------------------------------------------------|
| a1b2c3...       Chrome / Windows      192.168.1.10     2026-04-18      Active  |
| d4e5f6...       Edge / Android        192.168.1.11     2026-04-17      Active  |
|--------------------------------------------------------------------------------|
| [ Revoke Selected Session ]   [ Revoke Other Sessions ]                       |
+--------------------------------------------------------------------------------+
"""
    ascii_figure('Fig 7.5 Session Management Wireframe', session_wireframe)
    doc.add_page_break()

    # CHAPTER 8
    chapter_heading('CHAPTER 8: CHALLENGES FACED')
    section_heading('8.1 Cryptographic Design Challenges')
    body('Balancing usability and cryptographic hardness required careful parameter selection. High PBKDF2 iteration counts improve security but impact unlock latency on low-power devices. The project handled this through moderate UX buffering while retaining strong derivation parameters.')
    section_heading('8.2 Session and Security Middleware Integration')
    body('Combining JWT cookie auth, CSRF checks, rate limiting, and per-session database tracking required strict middleware ordering and consistent error handling. Integration testing was essential to avoid false negatives in authenticated routes.')
    section_heading('8.3 Risk Intelligence Reliability')
    body('Breach-check logic required robust fallback behavior for external API failures and careful privacy controls during hash prefix queries. The solution preserved user privacy by keeping full hash and plaintext comparison strictly client-side.')
    section_heading('8.4 Data Recovery and Rollback')
    body('Because plaintext data is never retained server-side, rollback had to be implemented with encrypted version snapshots only. This required version indexing and explicit rollback endpoints while preserving zero-knowledge constraints.')
    section_heading('8.5 Testing and Validation Constraints')
    body('Security-sensitive systems need both functionality checks and negative-path checks such as tamper validation, invalid token behavior, and rate-limit triggering. Designing these tests without weakening architecture assumptions was a significant engineering task.')
    doc.add_page_break()

    # CHAPTER 9
    chapter_heading('CHAPTER 9: CONCLUSION AND FUTURE SCOPE')
    section_heading('9.1 Conclusion')
    body('This project demonstrates a practical, full-stack, zero-knowledge password manager with integrated client-side risk intelligence. The implemented architecture ensures that cryptographic trust resides with the user, not the server. With modular design, hardened API controls, and user-centric security feedback, the system provides a strong academic and technical foundation for privacy-preserving credential management.')
    section_heading('9.2 Future Scope')
    bullet('WebAuthn-based biometric and hardware-key assisted authentication.')
    bullet('Secure encrypted sharing workflows using public-key cryptography between users.')
    bullet('Browser extension for autofill with strict local decryption boundaries.')
    bullet('Mobile-first client applications with synchronized encrypted vault access.')
    bullet('Risk analytics enhancements with adaptive recommendations and anomaly detection.')
    bullet('Production deployment hardening with managed certificate rotation and SIEM integration.')
    doc.add_page_break()

    # REFERENCES (IEEE, alphabetically arranged)
    chapter_heading('REFERENCES')
    refs = [
        '[1] Bitwarden, "Bitwarden Help Center," 2026. [Online]. Available: https://bitwarden.com/help/. [Accessed: Apr. 18, 2026].',
        '[2] M. Burnett and D. Kleiman, Perfect Passwords. Syngress, 2006.',
        '[3] M. Conti, N. Dragoni, and V. Lesyk, "A Survey of Man in the Middle Attacks," IEEE Commun. Surveys Tuts., vol. 18, no. 3, pp. 2027-2051, 2016.',
        '[4] MongoDB, "MongoDB Manual," 2026. [Online]. Available: https://www.mongodb.com/docs/manual/. [Accessed: Apr. 18, 2026].',
        '[5] NIST, "Digital Identity Guidelines: Authentication and Lifecycle Management (SP 800-63B)," 2023.',
        '[6] Node.js, "Node.js Documentation," 2026. [Online]. Available: https://nodejs.org/docs/latest/api/. [Accessed: Apr. 18, 2026].',
        '[7] OWASP Foundation, "Password Storage Cheat Sheet," 2026. [Online]. Available: https://cheatsheetseries.owasp.org/. [Accessed: Apr. 18, 2026].',
        '[8] React, "React Documentation," 2026. [Online]. Available: https://react.dev/. [Accessed: Apr. 18, 2026].',
        '[9] T. Hunt, "Have I Been Pwned Pwned Passwords API v3," 2026. [Online]. Available: https://haveibeenpwned.com/API/v3. [Accessed: Apr. 18, 2026].',
        '[10] W3C, "Web Cryptography API," 2017. [Online]. Available: https://www.w3.org/TR/WebCryptoAPI/. [Accessed: Apr. 18, 2026].',
    ]
    for ref in refs:
        body(ref)
    doc.add_page_break()

    # APPENDICES
    chapter_heading('APPENDICES')
    section_heading('APPENDIX A: API ENDPOINT SUMMARY')
    body('Auth: POST /api/auth/register, POST /api/auth/login, POST /api/auth/logout, GET /api/auth/me, POST /api/auth/mfa/setup, POST /api/auth/mfa/verify, POST /api/auth/mfa/disable, GET /api/auth/sessions, DELETE /api/auth/sessions/:id, GET /api/auth/audit-logs.')
    body('Vault: GET /api/vault, PUT /api/vault, GET /api/vault/versions, POST /api/vault/rollback.')
    section_heading('APPENDIX B: SECURITY CONFIGURATION SNAPSHOT')
    body('Security controls include Helmet headers, CSRF token validation, rate limiting for auth endpoints, JWT httpOnly cookies, input validation and sanitization, encrypted vault-only persistence, and audit event logging.')
    section_heading('APPENDIX C: MODULE-WISE FILE MAP')
    body('Frontend modules: auth.js, crypto.js, vault.js, risk.js, generator.js, app.js. Backend modules: server.js, app.js, routes/auth.js, routes/vault.js, middleware/auth.js, middleware/sanitize.js, models/User.js, models/Session.js, models/AuditLog.js, models/VaultVersion.js.')

    output_name = 'Final_Project_Report_Professional_Diagrams.docx'
    doc.save(output_name)
    print(f'Report generated successfully: {output_name}')


if __name__ == '__main__':
    create_report()
